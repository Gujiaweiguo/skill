"""表格样式系统 - 针对不同表格类型的样式渲染"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .parser import TableData


# ============================================================
# 公用辅助函数
# ============================================================

def _set_cell_shading(cell, color: str):
    """设置单元格底纹"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


def _set_cell_text(cell, text: str, bold: bool = False,
                   size: int = 10, alignment: int = 0,
                   color: str | None = None):
    """设置单元格文本样式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
    # 确保中文字体
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:ascii"), "Times New Roman")


def _apply_table_borders(table, sz: str = "4", color: str = "000000"):
    """给表格加边框"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_column_widths(table, widths_pct: list[int]):
    """设置列宽百分比"""
    if not widths_pct:
        return
    total_cols = len(table.columns)
    for i, width_pct in enumerate(widths_pct[:total_cols]):
        for row in table.rows:
            cell = row.cells[i]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tblW = table._tbl.find(qn("w:tblPr"))
            if tblW is not None:
                tblW_xml = tblW.find(qn("w:tblW"))
                if tblW_xml is not None:
                    total_w = int(tblW_xml.get(qn("w:w"), "9072"))
                    w = int(total_w * width_pct / 100)
                    tcW = parse_xml(
                        f'<w:tcW {nsdecls("w")} w:w="{w}" w:type="dxa"/>'
                    )
                    tcPr.append(tcW)


def _is_last_row(table, row_idx: int) -> bool:
    """判断是否为最后一行"""
    return row_idx == len(table.rows) - 1


# ============================================================
# 表格类型渲染函数
# ============================================================

def render_default_table(doc: Document, table_data: TableData):
    """通用表格 - 表头灰底+加粗，基础边框"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="4", color="999999")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "D9D9D9")

    # 数据行
    for ri, row_data in enumerate(table_data.rows):
        for ci, val in enumerate(row_data):
            _set_cell_text(table.rows[ri+1].cells[ci], val, size=10)

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")  # 表后间距


def render_comparison_table(doc: Document, table_data: TableData):
    """对比/响应表 - 左列加粗，交替行色"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="4", color="999999")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "2F5496")
        # 白色字
        tc = table.rows[0].cells[i]
        for p in tc.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)

    # 数据行
    for ri, row_data in enumerate(table_data.rows):
        for ci, val in enumerate(row_data):
            align = 3 if ci == len(row_data) - 1 else 0  # 最后一列居中
            _set_cell_text(table.rows[ri+1].cells[ci], val, size=10, alignment=align)
            # 左列加粗
            if ci == 0:
                for p in table.rows[ri+1].cells[ci].paragraphs:
                    for r in p.runs:
                        r.bold = True
            # 交替行色
            if ri % 2 == 1:
                _set_cell_shading(table.rows[ri+1].cells[ci], "F2F2F2")

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")


def render_function_matrix(doc: Document, table_data: TableData):
    """功能清单表 - 表头深蓝底纹+白字"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="6", color="2F5496")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "2F5496")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)

    # 数据行
    for ri, row_data in enumerate(table_data.rows):
        for ci, val in enumerate(row_data):
            align = 1 if ci == len(row_data) - 1 else 0  # 优先级列居中
            _set_cell_text(table.rows[ri+1].cells[ci], val, size=10, alignment=align)
            # 隔行灰色
            if ri % 2 == 1:
                _set_cell_shading(table.rows[ri+1].cells[ci], "F2F2F2")

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")


def render_pricing_table(doc: Document, table_data: TableData):
    """报价表 - 金额右对齐，合计行加粗+底纹"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="4", color="333333")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "D9D9D9")

    # 数据行
    for ri, row_data in enumerate(table_data.rows):
        is_total = row_data[0] in ("合计", "总计", "合计（含税）") or "合计" in row_data[0]
        for ci, val in enumerate(row_data):
            # 金额列右对齐（后两列）
            is_money_col = ci >= len(row_data) - 2
            align = 2 if is_money_col else 0
            _set_cell_text(table.rows[ri+1].cells[ci], val,
                           bold=is_total, size=10, alignment=align)
        # 合计行底纹
        if is_total:
            for ci in range(len(row_data)):
                _set_cell_shading(table.rows[ri+1].cells[ci], "F2F2F2")

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")


def render_implementation_plan(doc: Document, table_data: TableData):
    """实施计划表 - 阶段行底色标记"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="4", color="999999")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "2F5496")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)

    # 阶段颜色循环
    phase_colors = ["E8F0FE", "FFF2CC", "E8F5E9", "FCE4EC", "F3E5F5", "E0F7FA"]

    for ri, row_data in enumerate(table_data.rows):
        for ci, val in enumerate(row_data):
            _set_cell_text(table.rows[ri+1].cells[ci], val, size=10)
            # 阶段列（第一列）加颜色
            if ci == 0:
                color_idx = ri % len(phase_colors)
                _set_cell_shading(table.rows[ri+1].cells[ci], phase_colors[color_idx])

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")


def render_personnel_matrix(doc: Document, table_data: TableData):
    """人员配置表"""
    if not table_data.header or not table_data.rows:
        return

    rows = len(table_data.rows) + 1
    cols = len(table_data.header)
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_borders(table, sz="4", color="999999")

    # 表头
    for i, h in enumerate(table_data.header):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=1)
        _set_cell_shading(table.rows[0].cells[i], "2F5496")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)

    for ri, row_data in enumerate(table_data.rows):
        for ci, val in enumerate(row_data):
            _set_cell_text(table.rows[ri+1].cells[ci], val, size=10)
            # "待确认"标记红色
            if val == "待确认":
                for p in table.rows[ri+1].cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                        r.font.size = Pt(9)

    _set_column_widths(table, table_data.column_widths)
    doc.add_paragraph("")


# ============================================================
# 调度映射
# ============================================================

TABLE_RENDERERS = {
    "default-table": render_default_table,
    "comparison-table": render_comparison_table,
    "comparison": render_comparison_table,
    "function-matrix": render_function_matrix,
    "pricing-table": render_pricing_table,
    "implementation-plan": render_implementation_plan,
    "personnel-matrix": render_personnel_matrix,
}


def render_table(doc: Document, table_data: TableData):
    """调度到对应的表格渲染函数"""
    renderer = TABLE_RENDERERS.get(table_data.table_type, render_default_table)
    renderer(doc, table_data)