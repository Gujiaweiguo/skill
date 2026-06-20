"""word-master 文档渲染器

接收 ContentPackage → 加载模板 → 逐章渲染 → 输出 .docx
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .parser import ContentPackage, Chapter, TableData
from .table_styles import render_table


# ============================================================
# 路径变量解析
# ============================================================

ENV_VARS = {
    "$MATERIALS_DIR": os.environ.get("LANLNK_BASE", "/opt/code/docs/lanlnk") + "/materials",
    "$LANLNK_BASE": os.environ.get("LANLNK_BASE", "/opt/code/docs/lanlnk"),
}

WORD_TEMPLATES_DIR = Path(
    os.environ.get(
        "WORD_TEMPLATES_DIR",
        "/opt/code/skill/skills/word/word-master/templates",
    )
)


def resolve_path(path_str: str) -> str:
    """解析路径中的环境变量"""
    for var, val in ENV_VARS.items():
        path_str = path_str.replace(var, val)
    return path_str


# ============================================================
# 模板查找
# ============================================================

TEMPLATE_MAP = {
    "bidding-technical": "bidding-technical-base.docx",
    "bidding-commercial": "bidding-commercial-base.docx",
    "bidding-standard": "bidding-technical-base.docx",
    "bidding-compilation": "bidding-technical-base.docx",
    "proposal": "bidding-technical-base.docx",
    "report": "bidding-technical-base.docx",
    "intro": "bidding-technical-base.docx",
}


def find_template(template_name: str) -> Path | None:
    """查找模板文件"""
    filename = TEMPLATE_MAP.get(template_name)
    if not filename:
        return None
    tpl_path = WORD_TEMPLATES_DIR / filename
    return tpl_path if tpl_path.exists() else None


# ============================================================
# 渲染器
# ============================================================

class Renderer:
    """文档渲染器"""

    def __init__(self, pkg: ContentPackage):
        self.pkg = pkg

    def render(self, output_path: str | Path) -> Path:
        """渲染整个文档"""
        output_path = Path(output_path)

        # 1. 加载模板
        doc = self._load_template()

        # 2. 应用格式覆盖（从招标要求提取的字体/页边距/页面设置）
        if self.pkg.format_overrides:
            self._apply_format_overrides(doc)

        # 3. 构建封面
        self._build_cover(doc)

        # 3.5 清理模板占位正文（保留封面）
        self._clear_template_body(doc)

        # 4. 插入分页，开始正文
        doc.add_page_break()

        # 5. 预留目录页
        if self.pkg.toc and self.pkg.toc.enabled:
            self._build_toc(doc)

        # 6. 逐章渲染
        for chapter in self.pkg.chapters:
            self._render_chapter(doc, chapter)

        # 7. 设置页眉页脚
        self._apply_header_footer(doc)

        # 8. 保存
        doc.save(str(output_path))
        return output_path

    def _apply_format_overrides(self, doc: Document):
        """应用格式覆盖（响应招标格式要求）

        支持从 content package 的 format_overrides 字段中读取：
        - font:    body/heading 字体名称、字号
        - margins: top/bottom/left/right 页边距
        - page:    size (A4/A3/Letter)、orientation (portrait/landscape)
        """
        overrides = self.pkg.format_overrides
        if not overrides:
            return

        # ── 字体 ──
        if "font" in overrides:
            f = overrides["font"]
            normal = doc.styles["Normal"]
            rpr = normal.element.find(qn("w:rPr"))
            if rpr is None:
                rpr = parse_xml(f'<w:rPr {nsdecls("w")} />')
                normal.element.append(rpr)
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rFonts)

            # 正文东亚字体（如宋体）
            if "body" in f:
                rFonts.set(qn("w:eastAsia"), f["body"])
                rFonts.set(qn("w:ascii"), f.get("ascii", f["body"]))
                rFonts.set(qn("w:hAnsi"), f.get("ascii", f["body"]))

            # 正文字号（pt）
            if "size" in f:
                try:
                    normal.font.size = Pt(int(f["size"]))
                except (ValueError, TypeError):
                    pass

            # 标题字体
            heading_font = f.get("heading", f.get("body", "黑体"))
            for i in range(1, 5):
                style_name = f"Heading {i}"
                if style_name not in [s.name for s in doc.styles]:
                    continue
                hs = doc.styles[style_name]
                hrpr = hs.element.find(qn("w:rPr"))
                if hrpr is None:
                    hrpr = parse_xml(f'<w:rPr {nsdecls("w")} />')
                    hs.element.append(hrpr)
                hrFonts = hrpr.find(qn("w:rFonts"))
                if hrFonts is None:
                    hrFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                    hrpr.insert(0, hrFonts)
                hrFonts.set(qn("w:eastAsia"), heading_font)
                hrFonts.set(qn("w:ascii"), f.get("ascii", "Times New Roman"))
                hrFonts.set(qn("w:hAnsi"), f.get("ascii", "Times New Roman"))

                # 标题字号（可选）
                heading_size_key = f"heading_{i}_size"
                if heading_size_key in f:
                    try:
                        hs.font.size = Pt(int(f[heading_size_key]))
                    except (ValueError, TypeError):
                        pass

        # ── 页边距 ──
        if "margins" in overrides:
            m = overrides["margins"]
            for section in doc.sections:
                if "top" in m:
                    try:
                        section.top_margin = Cm(float(m["top"]))
                    except (ValueError, TypeError):
                        pass
                if "bottom" in m:
                    try:
                        section.bottom_margin = Cm(float(m["bottom"]))
                    except (ValueError, TypeError):
                        pass
                if "left" in m:
                    try:
                        section.left_margin = Cm(float(m["left"]))
                    except (ValueError, TypeError):
                        pass
                if "right" in m:
                    try:
                        section.right_margin = Cm(float(m["right"]))
                    except (ValueError, TypeError):
                        pass

        # ── 页面大小 ──
        if "page" in overrides:
            p = overrides["page"]
            for section in doc.sections:
                # 纸张大小
                size_name = p.get("size", "A4").upper()
                if size_name == "A3":
                    section.page_width = Cm(42)
                    section.page_height = Cm(29.7)
                elif size_name == "A4":
                    section.page_width = Cm(21)
                    section.page_height = Cm(29.7)
                elif size_name == "LETTER":
                    section.page_width = Cm(21.59)
                    section.page_height = Cm(27.94)

                # 方向
                orientation = p.get("orientation", "portrait")
                if orientation == "landscape":
                    section.page_width, section.page_height = (
                        section.page_height, section.page_width
                    )

        # ── 行距 ──
        if "line_spacing" in overrides:
            try:
                spacing_val = float(overrides["line_spacing"])
                for p in doc.paragraphs:
                    p.paragraph_format.line_spacing = spacing_val
            except (ValueError, TypeError):
                pass

    def _load_template(self) -> Document:
        """加载模板文件"""
        # 根据 type 或 template 选择模板
        template_name = self.pkg.template or self.pkg.type
        tpl_path = find_template(template_name)

        if tpl_path:
            return Document(str(tpl_path))

        # 无模板 → 新建空白文档
        doc = Document()

        # 设置默认 A4
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # Normal
        normal = doc.styles['Normal']
        normal.font.size = Pt(12)
        rpr = normal.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = parse_xml(f'<w:rPr {nsdecls("w")} />')
            normal.element.append(rpr)
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")

        return doc

    def _build_cover(self, doc: Document):
        """构建封面 — 替换模板中的占位文本"""
        cv = self.pkg.cover
        if not cv:
            return

        found_title = False
        found_subtitle = False
        found_date = False

        for p in doc.paragraphs:
            if "[项目名称]" in p.text and not found_title:
                for r in p.runs:
                    if "[项目名称]" in r.text:
                        r.text = r.text.replace("[项目名称]", cv.title)
                        found_title = True

        subtitle = cv.subtitle
        if not subtitle:
            subtitle = self._default_subtitle()
        if subtitle and not found_subtitle:
            for p in doc.paragraphs:
                if "技术方案" in p.text or "商务标" in p.text or "技术标" in p.text:
                    for r in p.runs:
                        if r.text in ("技术方案", "商务标", "技术标"):
                            r.text = subtitle
                            found_subtitle = True

        if cv.date and not found_date:
            for p in doc.paragraphs:
                if "年" in p.text and "月" in p.text and len(p.text.strip()) <= 15:
                    for r in p.runs:
                        if r.text.strip():
                            r.text = cv.date
                            found_date = True
        elif self.pkg.date and not found_date:
            for p in doc.paragraphs:
                if "年" in p.text and "月" in p.text and len(p.text.strip()) <= 15:
                    for r in p.runs:
                        if r.text.strip():
                            r.text = self.pkg.date
                            found_date = True

    def _default_subtitle(self) -> str:
        """根据文档类型返回默认副标题"""
        defaults = {
            "bidding-technical": "技术方案",
            "bidding-commercial": "商务标",
            "proposal": "方案建议书",
            "report": "立项报告",
            "intro": "公司介绍",
        }
        return defaults.get(self.pkg.template or self.pkg.type, "技术方案")

    def _clear_template_body(self, doc: Document):
        """清理模板中封面之后的占位正文段落（如"第一章 综述""此处为正文内容"）"""
        body = doc.element.body
        # 策略：删除所有含文本的 Heading 段落 + 其后紧跟的占位文本段落
        # 封面段落都是 Normal 样式且不含 Heading 格式
        to_remove = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            text = p.text.strip()
            if not text:
                continue
            # 封面文字（项目名/技术方案/公司名/日期）不动
            # 占位正文（Heading 样式 + "此处为正文" 等）删除
            if "Heading" in style_name:
                to_remove.append(p)
            elif text.startswith("此处为正文") or text.startswith("（正文使用"):
                to_remove.append(p)

        for p in to_remove:
            p._element.getparent().remove(p._element)

    def _build_toc(self, doc: Document):
        """构建目录页占位"""
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run("目  录")
        run.font.size = Pt(16)
        run.font.bold = True
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "黑体")
        rFonts.set(qn("w:ascii"), "微软雅黑")

        # 列出标题
        for chapter in self.pkg.chapters:
            self._add_toc_entry(doc, chapter.heading_text, 0, chapter)
            for sub in chapter.sub_chapters:
                if self.pkg.toc and sub.heading_level <= self.pkg.toc.max_level:
                    self._add_toc_entry(doc, sub.heading_text, 1, sub)

        doc.add_page_break()

    def _add_toc_entry(self, doc: Document, text: str, indent: int,
                       chapter: Chapter):
        """添加目录条目"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6) if indent == 0 else Pt(2)
        p.paragraph_format.space_after = Pt(2)

        indent_text = "    " * indent + text
        run = p.add_run(indent_text)
        run.font.size = Pt(12) if indent == 0 else Pt(11)
        if indent == 0:
            run.font.bold = True
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")

    def _render_chapter(self, doc: Document, chapter: Chapter):
        """渲染单个章节"""
        # 分页
        if chapter.page_break and doc.paragraphs:
            doc.add_page_break()

        # 标题
        style_name = f"Heading {chapter.heading_level}"
        if style_name in [s.name for s in doc.styles]:
            heading = doc.add_paragraph(chapter.heading_text, style=style_name)
        else:
            heading = doc.add_paragraph(chapter.heading_text)
            heading.style = doc.styles['Heading 1']

        # 渲染段落
        self._render_paragraphs(doc, chapter.paragraphs)

        # 渲染无序列表
        self._render_list_items(doc, chapter.list_items)

        # 渲染有序列表
        self._render_ordered_list(doc, chapter.ordered_list_items)

        # 渲染表格
        for table in chapter.tables:
            render_table(doc, table)

        # 渲染图片
        for img in chapter.images:
            self._render_image(doc, img)

        # 渲染子章节
        for sub in chapter.sub_chapters:
            self._render_chapter(doc, sub)

    def _render_paragraphs(self, doc: Document, paragraphs: list[str]):
        """渲染正文段落"""
        for text in paragraphs:
            # 检查是否是列表
            stripped = text.strip()
            if not stripped:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

            # 检测是否为引用块
            if stripped.startswith(">"):
                stripped = stripped.lstrip("> ").strip()
                p.paragraph_format.left_indent = Cm(1.0)
                # 灰色左竖线
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="999999"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)

            run = p.add_run(stripped)
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "Times New Roman")

    def _render_list_items(self, doc: Document, items: list[list[str]]):
        """渲染无序列表"""
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(2)

            # 项目符号
            run_bullet = p.add_run("● ")
            run_bullet.font.size = Pt(8)
            run_bullet.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            text = item[0] if isinstance(item, list) else item
            run = p.add_run(text)
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")

    def _render_ordered_list(self, doc: Document, items: list[list[str]]):
        """渲染有序列表"""
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(2)

            run_num = p.add_run(f"{idx}. ")
            run_num.font.size = Pt(12)

            text = item[0] if isinstance(item, list) else item
            run = p.add_run(text)
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")

    def _render_image(self, doc: Document, img: dict):
        """渲染图片"""
        img_path = resolve_path(img.get("path", ""))
        alt_text = img.get("alt", "")

        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = 1  # 居中
                run = p.add_run()
                run.add_picture(img_path)
                # 图注
                if alt_text:
                    cap = doc.add_paragraph()
                    cap.alignment = 1
                    r = cap.add_run(alt_text)
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            except Exception as e:
                p = doc.add_paragraph()
                run = p.add_run(f"⚠️ 图片加载失败: {img_path} ({e})")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ 图片未找到: {img_path}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

    def _apply_header_footer(self, doc: Document):
        """设置页眉页脚"""
        pkg = self.pkg

        # 页眉
        if pkg.header:
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False
                hp = header.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 左：公司名
                left_text = pkg.header.left or pkg.author or ""
                if left_text:
                    run = hp.add_run(left_text)
                    run.font.size = Pt(9)

        # 页脚
        if pkg.footer:
            for section in doc.sections:
                footer = section.footer
                footer.is_linked_to_previous = False
                fp = footer.paragraphs[0]
                fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                right_text = pkg.footer.right or "第 {page} 页"
                # 替换 {page} 为页码域
                parts = re.split(r"\{page\}", right_text)
                for i, part in enumerate(parts):
                    if part:
                        run = fp.add_run(part)
                        run.font.size = Pt(9)
                    if i < len(parts) - 1:
                        # 插入页码域
                        fldChar = parse_xml(
                            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
                        )
                        run_f = fp.add_run()
                        run_f._element.append(fldChar)
                        instrText = parse_xml(
                            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
                        )
                        run_i = fp.add_run()
                        run_i._element.append(instrText)
                        fldChar2 = parse_xml(
                            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
                        )
                        run_e = fp.add_run()
                        run_e._element.append(fldChar2)


def render(pkg: ContentPackage, output_path: str | Path) -> Path:
    """便捷函数：渲染文档"""
    return Renderer(pkg).render(output_path)